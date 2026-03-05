from pathlib import Path

from docx import Document


def parse(path: Path, max_pages: int | None = None) -> str:
    """Parse DOCX files and extract text from paragraphs and tables."""
    doc = Document(str(path))
    parts: list[str] = []

    # Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Extract tables
    for i, table in enumerate(doc.tables, 1):
        rows: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("\t".join(cells))
        if rows:
            parts.append(f"\n[Table {i}]\n" + "\n".join(rows))

    if not parts:
        return "(No text content could be extracted from this DOCX)"

    return "\n\n".join(parts)
